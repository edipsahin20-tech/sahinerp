from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name("SahinSoft_Altyapi_Denetim_Raporu.docx")
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT = "E8EEF5"
PALE = "F2F4F7"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 14, 7),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 11.5, DARK, 7, 3),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ["List Bullet", "List Number"]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.42)
    st.paragraph_format.first_line_indent = Inches(-0.2)
    st.paragraph_format.space_after = Pt(3)

header = sec.header.paragraphs[0]
header.text = "ŞAHİNSOFT  |  Altyapı Denetim Raporu"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("ŞahinSoft Ön Muhasebe ve Stok Yönetim Sistemi  •  27 Temmuz 2026  •  ")
run.font.size = Pt(8)
run.font.color.rgb = GRAY
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
footer._p.append(fld)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def table(headers, rows, widths, header_fill=LIGHT):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[i]
        c.width = Inches(w)
        c.text = h
        shade(c, header_fill)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c)
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = DARK
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, (value, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w)
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.size = Pt(8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def status_table(rows):
    t = table(["Durum", "Konu", "Denetim sonucu"], rows, [0.9, 1.8, 3.95])
    for row in t.rows[1:]:
        val = row.cells[0].text
        shade(row.cells[0], GREEN if val == "Hazır" else AMBER if val == "Kısmi" else RED)
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.68)
    p.add_run(text)
    return p

current_numbering_id = None

def start_numbered_list():
    global current_numbering_id
    numbering = doc.part.numbering_part.element
    abstract_ids = numbering.xpath("./w:abstractNum/@w:abstractNumId")
    abstract_id = abstract_ids[0] if abstract_ids else "0"
    existing_ids = [int(x) for x in numbering.xpath("./w:num/@w:numId")]
    current_numbering_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(current_numbering_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    if current_numbering_id is not None:
        pPr = p._p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(current_numbering_id))
        numPr.append(ilvl)
        numPr.append(numId)
    p.add_run(text)
    return p

def callout(title, text, fill=PALE):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Inches(6.65)
    shade(c, fill)
    set_cell_margins(c, 140, 180, 140, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = DARK
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_page_break():
    doc.add_page_break()

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("TEKNİK ALTYAPI DENETİM RAPORU")
r.font.name = "Calibri"
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("ŞahinSoft Ön Muhasebe, Stok, Cari ve Teklif Yönetim Sistemi")
r.font.size = Pt(14)
r.font.color.rgb = BLUE
r.font.bold = True

table(
    ["Alan", "Bilgi"],
    [
        ("İncelenen çözüm", "ASP.NET Core 10 / EF Core 10 / SQL Server 2022"),
        ("Denetim kapsamı", "Entity, ilişki, indeks, constraint, migration, servis, güvenlik ve modül olgunluğu"),
        ("Denetim tarihi", "27 Temmuz 2026"),
        ("Sonuç", "Temel veri modeli güçlü; canlı kullanımdan önce kritik altyapı düzeltmeleri ve uçtan uca testler gerekli"),
    ],
    [1.55, 5.1],
)

callout(
    "Raporun amacı",
    "Bu belge yalnızca yapılanları anlatmaz. Hangi altyapının gerçekten çalışır durumda olduğunu, hangisinin yalnızca veritabanı modeli olarak bulunduğunu ve hangi risklerin giderilmeden canlıya çıkılmaması gerektiğini açıkça ayırır.",
    LIGHT,
)

doc.add_heading("Yönetici özeti", level=1)
doc.add_paragraph(
    "Çözüm; ürün, kategori, barkod, varyant, depo, stok hareketi, cari, teklif, alış/satış faturası, fiyat listesi, tahsilat/ödeme, kasa-banka, sipariş, irsaliye, masraf, çek-senet, sayım ve entegrasyon kayıtlarını kapsayan geniş bir ilişkisel modele sahiptir. Son derleme 0 hata ve 0 uyarı ile tamamlanmış, EF Core modeli ile son migration arasında bekleyen değişiklik bulunmadığı doğrulanmıştır."
)
doc.add_paragraph(
    "Bununla birlikte sistem henüz tamamlanmış bir ön muhasebe uygulaması değildir. Ürün kartı dışındaki modüllerin çoğu için kullanıcı ekranı ve controller yoktur. Bazı posting servisleri yazılmış olsa da iptal/ters kayıt, tahsilat posting, açılış stoğu, audit üretimi ve otomatik test altyapısı tamamlanmamıştır."
)

status_table([
    ("Hazır", "Çözüm ve derleme", ".NET 10 çözümü derleniyor; 0 hata, 0 uyarı."),
    ("Hazır", "EF Core model uyumu", "Son migration ile model uyumlu; pending model change yok."),
    ("Hazır", "Temel ilişkisel şema", "Ana tablolar, foreign key’ler, indeksler ve pek çok check constraint tanımlı."),
    ("Kısmi", "Stok posting", "Fatura, transfer, stok fişi ve sayım servisleri var; ekran, iptal ve tam tutarlılık testleri eksik."),
    ("Kısmi", "Cari ve finans", "Fatura ve tahsilat/ödeme posting servisleri var; ekran ve ters kayıt tamamlanmalı."),
    ("Kısmi", "Kimlik ve yetki", "Identity ve authentication hattı var; rol bazlı ekran politikaları tamamlanmalı."),
    ("Eksik", "Otomatik test", "Unit, integration ve SQL Server transaction testi bulunmuyor."),
    ("Eksik", "Canlı veritabanı uygulaması", "Son migration sunucudaki SQL Server’a henüz uygulanmadı."),
])

doc.add_heading("1. Denetlenen teknik yapı", level=1)
bullet("Hedef çerçeve: .NET 10.0.")
bullet("Web: ASP.NET Core MVC, Razor ve Bootstrap 5.")
bullet("Kimlik: ASP.NET Core Identity; Yönetici ve Personel rolleri.")
bullet("Veri erişimi: Entity Framework Core 10.0.10 ve SQL Server sağlayıcısı.")
bullet("Veritabanı hedefi: SQL Server 2022 Express, SahinSoftDb.")
bullet("Çözüm projeleri: Domain, Application, Infrastructure ve Web.")
bullet("Migration zinciri: Identity’den katalog, iş çekirdeği, finans, gelişmiş stok, barkod politikaları ve bütünleşik ön muhasebe çekirdeğine uzanıyor.")
bullet("Web sitesi kataloğundan 8 kategori ve 47 ürün için başlangıç verisi hazırlandı.")
bullet("KDV tanımları: %0, %1, %10 ve %20.")

doc.add_heading("2. Katman mimarisi denetimi", level=1)
table(
    ["Katman", "Mevcut durum", "Değerlendirme"],
    [
        ("Domain", "Entity ve enumlar burada.", "Doğru başlangıç; iş kurallarının bir kısmı halen Web servislerinde."),
        ("Application", "Proje var, içerik büyük ölçüde boş.", "Use-case arayüzleri, DTO ve validasyonlar taşınmalı."),
        ("Infrastructure", "Proje var, içerik büyük ölçüde boş.", "DbContext, repository/entegrasyon ve dış servis uygulamaları burada olmalı."),
        ("Web", "DbContext, migration, posting servisleri, controller ve UI burada.", "Şu an çalışan merkez; katmanlı yapı fiziksel olarak var fakat sorumluluk ayrımı tamamlanmamış."),
    ],
    [1.05, 2.45, 3.15],
)
callout(
    "Altyapı kararı",
    "Mevcut yapı prototip geliştirme için çalışır; büyüme hedefi açısından posting servisleri ve DbContext bağımlılıkları Application/Infrastructure katmanlarına ayrılmalıdır.",
    AMBER,
)

doc.add_heading("3. Ortak kayıt kimliği ve izlenebilirlik", level=1)
doc.add_paragraph(
    "Domain’de EntityBase kullanan iş tablolarına hem sayısal Id hem de RecordId (GUID) verilmektedir. SQL Server tarafında RecordId için NEWSEQUENTIALID() varsayılanı ve benzersiz indeks tanımlanmıştır."
)
bullet("Id: hızlı iç ilişki ve clustered primary key amacıyla.")
bullet("RecordId: dış sistem, API, senkronizasyon ve güvenli genel tanımlayıcı amacıyla.")
bullet("CreatedAtUtc / UpdatedAtUtc: kayıt zaman takibi amacıyla.")
bullet("ExternalRecordMapping: Dinosoft, Mikro veya başka kaynakların dış kimliklerini iç kayda bağlamak için.")
bullet("IntegrationOutboxMessage: onaylanan olayların güvenilir entegrasyon kuyruğuna yazılması için.")
callout(
    "Sınır",
    "GUID altyapısı iş tablolarını kapsar; ASP.NET Identity kullanıcı tabloları kendi string kimlik yapısını kullanır. Bu normaldir.",
)

doc.add_heading("4. Ana tablo grupları ve bağlantıları", level=1)
table(
    ["Grup", "Başlıca tablolar", "Bağlantı mantığı"],
    [
        ("Stok kartı", "Products, Categories, TaxRates, Barcodes, Variants, Colors, Images", "Ürün; kategori ve KDV’ye zorunlu, varyant/barkod/görsele çoklu bağlı."),
        ("Depo", "Branches, Warehouses, StockMovements, Transfers", "Şube → depo; tüm miktar değişimleri depo bazlı hareket üzerinden izlenir."),
        ("Ticari belge", "Quotes, Orders, DispatchNotes, Invoices", "Tekliften siparişe, irsaliyeye ve faturaya referans taşınabilir."),
        ("Cari", "Customers, Addresses, Contacts, AccountTransactions", "Belge hareketleri müşteri hesabının borç/alacak föyüne bağlanır."),
        ("Finans", "FinancialAccounts, Transactions, Receipts, Expenses", "Kasa/banka hareketleri cari ve belge hareketlerine bağlanabilir."),
        ("Fiyat", "PurchasePriceLists, SalesPriceLists ve satırları", "Cari, ürün, varyant, miktar ve geçerlilik tarihine göre fiyat altyapısı."),
        ("Kontrol", "AuditLogs, NumberSequences, Settings, Outbox, ExternalMappings", "Parametre, numara, iz ve entegrasyon güvenliği."),
    ],
    [1.1, 2.75, 2.8],
)

add_page_break()
doc.add_heading("5. Stok altyapısının satır satır iş mantığı", level=1)

doc.add_heading("5.1 Stok kartı", level=2)
bullet("Stok kodu veritabanında zorunlu ve benzersizdir.")
bullet("Kullanıcı kodu boş bırakırsa NumberSequence üzerinden SHN.001 biçiminde sıra üretilir.")
bullet("Ana barkod Products tablosunda, sınırsız ek barkod ProductBarcodes tablosunda tutulur.")
bullet("ProductBarcodes.Barcode benzersiz indekslidir; aynı barkod iki üründe kullanılamaz.")
bullet("Kategori ve KDV foreign key’leri Restrict davranışındadır; kullanılan tanımlar yanlışlıkla silinemez.")
bullet("Fiyatlar negatif olamaz; minimum stok negatif olamaz.")
bullet("Renk ve diğer seçenekler ProductVariant üzerinden ürüne bağlıdır.")
bullet("Ürün görselleri ürün veya varyanta bağlanabilir; birincil görsel işareti vardır.")
bullet("Seri ve lot takibi için ürün bayrakları ve ProductSerialNumbers tablosu bulunur.")

doc.add_heading("5.2 Barkod", level=2)
bullet("EAN-13 ve EAN-8 için kontrol hanesi hesaplayan servis vardır.")
bullet("Terazi barkodu 27, 28 veya 29 ön eki + 5 haneli PLU olacak şekilde 7 hanedir.")
bullet("Terazi ürünü için Kg/Adet ölçüm türü ve fiyat içerme seçeneği saklanır.")
bullet("Barkod zorunluluğu ve otomatik üretim InventorySettings üzerinden parametriktir.")
callout(
    "Denetim sırasında düzeltildi",
    "Ürün düzenleme POST işleminde mevcut barkod koleksiyonu artık yüklenerek eşitleniyor. Aynı ana barkodun ikinci kez eklenmesine yol açan hata kapatıldı.",
    GREEN,
)

doc.add_heading("5.3 Stok hareket defteri", level=2)
doc.add_paragraph(
    "StockMovements miktarın gerçek muhasebe defteri olacak şekilde tasarlanmıştır. Hareket; ürün, depo, tarih, miktar, maliyet, belge numarası ve kaynak satır kimliğini taşır."
)
table(
    ["Kaynak", "Hareket", "Bağlantı alanı"],
    [
        ("Alış faturası", "Pozitif Purchase", "InvoiceLineId"),
        ("Satış faturası", "Negatif Sale", "InvoiceLineId"),
        ("Depolar arası sevk", "Kaynakta TransferOut, hedefte TransferIn", "StockTransferLineId"),
        ("Stok giriş fişi", "Pozitif AdjustmentIn", "StockSlipLineId"),
        ("Stok çıkış fişi", "Negatif AdjustmentOut", "StockSlipLineId"),
        ("Sayım fazlası", "Pozitif InventoryCountSurplus", "InventoryCountLineId"),
        ("Sayım noksanı", "Negatif InventoryCountShortage", "InventoryCountLineId"),
    ],
    [1.35, 3.15, 2.15],
)

doc.add_heading("5.4 Sayım örneği", level=2)
start_numbered_list()
numbered("Sayım onay anında ilgili ürün ve depo için stok hareketleri toplanır.")
numbered("Sistem miktarı 5, sayılan miktar 4 ise fark 4 − 5 = −1 hesaplanır.")
numbered("Sayım satırına Sistem: 5, Sayılan: 4 bilgileri yazılır.")
numbered("Stok hareketine −1 miktarlı Sayım Noksanı kaydı atılır.")
numbered("Açıklama sistem, sayılan ve fark değerlerini içerir.")
numbered("Ürün toplam stoğu 4 olarak güncellenir.")
numbered("Tüm işlem tek transaction içinde kaydedilir; hata olursa hiçbir parça kalıcı olmaz.")

callout(
    "Denetim sırasında düzeltildi",
    "Yeni stok kartındaki başlangıç miktarı varsayılan depoya Açılış hareketi üretiyor. Düzenleme ekranında mevcut miktar salt okunur; sonraki değişimler fatura, fiş, sevk veya sayım üzerinden yapılmak zorunda.",
    GREEN,
)

add_page_break()
doc.add_heading("6. Fatura, cari ve finans bağlantısı", level=1)
doc.add_heading("6.1 Alış faturası", level=2)
start_numbered_list()
numbered("Taslak fatura ve satırları yüklenir.")
numbered("Miktar, fiyat, iskonto ve KDV tekrar hesaplanır.")
numbered("Stok takipli her ürün için depo stok hareketi pozitif yazılır.")
numbered("Ürünün toplam stok önbelleği artırılır.")
numbered("Cari hesaba alış kaynaklı alacak hareketi yazılır.")
numbered("Vade satırı yoksa tek vade otomatik oluşturulur.")
numbered("PurchaseInvoiceApproved olayı outbox’a yazılır.")
numbered("Fatura Approved olur ve tamamı tek Serializable transaction’da kaydedilir.")

doc.add_heading("6.2 Satış faturası", level=2)
start_numbered_list()
numbered("Taslak ve satırlar doğrulanır; toplamlar sunucuda yeniden hesaplanır.")
numbered("Stok parametreleri açıksa kullanılabilir stok ve rezervasyon kontrol edilir.")
numbered("Stok takipli ürün için negatif Sale hareketi oluşturulur.")
numbered("Cari hesaba satış kaynaklı borç hareketi yazılır.")
numbered("Vade ve entegrasyon olayı oluşturulur; fatura Approved olur.")

doc.add_heading("6.3 Borç–alacak yönü", level=2)
table(
    ["İşlem", "Cari borç", "Cari alacak", "Stok"],
    [
        ("Satış faturası", "Fatura toplamı", "0", "Azalır"),
        ("Alış faturası", "0", "Fatura toplamı", "Artar"),
        ("Tahsilat", "0", "Tahsilat tutarı", "Etkisiz"),
        ("Tediye/ödeme", "Ödeme tutarı", "0", "Etkisiz"),
    ],
    [1.7, 1.55, 1.55, 1.85],
)
doc.add_paragraph(
    "Dört işlem yönü için posting altyapısı bulunmaktadır. Tahsilat/tediye servisi cari ve kasa/banka hareketlerini aynı Serializable transaction içinde üretir; kullanıcı ekranı ve ters kayıt akışı sonraki adımdır."
)

doc.add_heading("6.4 İptal ve ters kayıt", level=2)
callout(
    "Canlıya çıkış engeli",
    "Onaylanmış fatura, stok fişi, sayım ve sevk için iptal/ters hareket servisi henüz yoktur. Ön muhasebede onaylı kayıt silinmemeli; aynı kaynak belgeye bağlı ters stok, ters cari ve ters finans hareketi üretilmelidir.",
    RED,
)

doc.add_heading("7. Parametrik yapı", level=1)
table(
    ["Parametre", "Mevcut amaç", "Denetim"],
    [
        ("RequireBarcode", "Barkod zorunluluğu", "Serviste kullanılıyor."),
        ("AutoGenerateBarcode", "Boş barkodu üretme", "Serviste kullanılıyor."),
        ("DefaultBarcodeType", "EAN13/EAN8 seçimi", "Ürün oluşturma akışında kullanılıyor."),
        ("EnforceStockLevel", "Stok seviyesi kontrolü", "Fatura, transfer ve stok çıkışında kullanılıyor."),
        ("AllowNegativeStock", "Negatif stok izni", "Posting kontrollerinde kullanılıyor."),
        ("AllowSaleWhenOutOfStock", "Stoksuz satış", "Satış faturasında kullanılıyor."),
        ("RequireTransferApproval", "Sevk onayı", "Modelde var; servis akışına tam politika olarak bağlanmalı."),
        ("TrackStockByVariant", "Varyant bazlı stok", "Model var; fatura kontrolü varyantı tam filtrelemiyor."),
        ("RequireProductVariant", "Varyant seçimini zorunlu yapma", "Varsayılan kapalı; açılırsa varyantlı üründe seçim zorunlu."),
        ("PreventSaleBelowCost", "Maliyet altı satış engeli", "Parametre var; henüz posting kuralına bağlanmamış."),
    ],
    [1.8, 2.45, 2.4],
)

add_page_break()
doc.add_heading("8. Veritabanı bütünlüğü denetimi", level=1)
doc.add_heading("8.1 Güçlü yönler", level=2)
bullet("Stok kodu, barkod, kategori kodu, depo kodu, belge numaraları ve birçok referans için benzersiz indeksler var.")
bullet("Para ve miktar alanları için açık decimal precision tanımları var.")
bullet("Negatif fiyat, sıfır miktarlı satır, aynı depo arası sevk ve aynı satır numarası gibi hataları engelleyen check/unique kuralları var.")
bullet("Master kayıt silmelerinde Restrict; belge satırlarında uygun yerlerde Cascade; opsiyonel kaynak bağlantılarında SetNull kullanılmış.")
bullet("Cari hareketi aynı anda hem borç hem alacak olamayacak biçimde check constraint ile korunuyor.")
bullet("Her iş tablosunda benzersiz RecordId indeksi var.")
bullet("Arama ve föy sorguları için ürün/depo/tarih, cari/tarih, belge numarası ve durum indeksleri var.")

doc.add_heading("8.2 Düzeltilmesi gereken veri kuralları", level=2)
table(
    ["Öncelik", "Bulgu", "Önerilen çözüm"],
    [
        ("Tamam", "Stok kartı açılışı ve manuel miktar riski.", "Açılış hareketi eklendi; düzenlemede miktar salt okunur."),
        ("P0", "Onaylı belgelerin iptal/ters kayıt altyapısı yok.", "Kaynak belge + ReversalOfRecordId ile ters hareket servisi kurulsun."),
        ("Tamam", "Authentication middleware sırası.", "UseAuthentication(), UseAuthorization() öncesine eklendi."),
        ("P0", "Otomatik test projesi yok.", "SQL Server tabanlı integration testleri ve hesaplama unit testleri eklensin."),
        ("Tamam", "Fatura stok kontrolünde varyant bağlantısı.", "InvoiceLine.ProductVariantId ve foreign key eklendi."),
        ("Tamam", "Ortak kullanılabilir stok hesabı.", "Depo + ürün + varyant + aktif rezervasyon ortak servise taşındı."),
        ("P1", "Sayım/stok fişi ekranı ve API/controller yok.", "Taslak, onay ve sonuç föyü uçtan uca tamamlanmalı."),
        ("Tamam", "Tahsilat/ödeme posting bağlantısı.", "Cari ve kasa/banka hareketleri tek transaction’da üretiliyor."),
        ("P1", "AuditLog tablosu var fakat otomatik kayıt üretimi yok.", "SaveChanges interceptor veya domain event ile kullanıcı/IP/değişiklik kaydı yazılmalı."),
        ("P2", "Application ve Infrastructure katmanları boş.", "İş servisleri ve EF uygulaması doğru katmanlara ayrılmalı."),
        ("P2", "QuestPDF paketi/servisi henüz yok.", "Teklif PDF modülünde lisans, font ve logo varlıklarıyla eklenmeli."),
    ],
    [0.65, 3.05, 2.95],
)

doc.add_heading("9. Güvenlik denetimi", level=1)
bullet("Identity parola kuralı güçlü: en az 10 karakter, büyük/küçük harf, rakam ve özel karakter.")
bullet("5 başarısız girişte 15 dakika kilitleme tanımlı.")
bullet("E-posta benzersizliği zorunlu.")
bullet("Administrator ve Staff rolleri seed ediliyor.")
bullet("Ürün controller’ı [Authorize] ile giriş gerektiriyor.")
bullet("Form POST işlemlerinde antiforgery doğrulaması var.")
callout(
    "Denetim sırasında düzeltildi",
    "Program.cs içinde app.UseAuthentication() çağrısı UseAuthorization() öncesine eklendi.",
    GREEN,
)
bullet("Rol bazlı yetkilendirme henüz ekran ve eylem seviyesinde uygulanmamış; ürün kartında tüm giriş yapan kullanıcılar aynı hakka sahip.")
bullet("Bootstrap yönetici parolası appsettings içinde tutulmuyor; ortam değişkeni/Secret Store kullanılması için uygun.")
bullet("Connection string Windows Authentication kullanıyor; IIS App Pool kimliğine SQL yetkisi verilmesi gerekecek.")

add_page_break()
doc.add_heading("10. Derleme ve migration doğrulaması", level=1)
table(
    ["Kontrol", "Sonuç", "Açıklama"],
    [
        ("Solution build", "Başarılı", "4 proje derlendi; 0 hata, 0 uyarı."),
        ("EF pending changes", "Başarılı", "Son migration sonrası model değişikliği yok."),
        ("Migration script", "Başarılı", "Son altyapı migration’ı SQL scriptine üretilebildi."),
        ("Test keşfi", "Başarısız/eksik", "Çözümde test projesi bulunmadı."),
        ("Sunucu DB uygulaması", "Yapılmadı", "Migration SQL Server 2022 Express’e henüz uygulanmadı."),
        ("Canlı IIS testi", "Yapılmadı", "Publish, App Pool, domain ve SSL daha sonraki aşama."),
    ],
    [1.7, 1.35, 3.6],
)

doc.add_heading("11. Modül olgunluk matrisi", level=1)
status_table([
    ("Hazır", "Ürün listeleme/kart", "Liste, arama, oluşturma ve düzenleme ekranı mevcut; barkod edit hatası düzeltilmeli."),
    ("Hazır", "Katalog seed", "8 kategori, 47 ürün ve 4 KDV oranı migration başlangıç verisinde."),
    ("Kısmi", "Barkod/varyant/görsel", "Şema geniş; ana stok kartı ekranı yalnızca birincil barkodu kullanıyor."),
    ("Kısmi", "Depo ve sevk", "Model ve onay servisi var; ekran/controller yok."),
    ("Kısmi", "Sayım", "Model ve fark posting servisi var; sayım giriş/sonuç ekranı yok."),
    ("Kısmi", "Alış/satış faturası", "Model ve onay posting servisi var; CRUD ekranı ve iptal yok."),
    ("Kısmi", "Cari", "Model ve fatura hareketi var; cari kart/föy ekranı yok."),
    ("Kısmi", "Tahsilat/tediye", "Posting servisi var; CRUD ekranı ve iptal/ters kayıt yok."),
    ("Kısmi", "Teklif", "Model var; hesaplama ekranı ve durum akışı yok."),
    ("Eksik", "PDF", "QuestPDF paketi, ŞahinSoft şablonu ve çıktı servisi yok."),
    ("Kısmi", "Sipariş/irsaliye/masraf/çek-senet", "İlişkisel şema var; iş servisleri ve ekranlar yok."),
    ("Eksik", "Raporlama/yedekleme", "Operasyonel raporlar ve otomatik yedek yapılandırması yapılmadı."),
])

doc.add_heading("12. Önerilen altyapı geliştirme sırası", level=1)
start_numbered_list()
numbered("P0: Stok açılışı ve authentication düzeltmeleri tamamlandı; bunlar için otomatik regression testleri yazın.")
numbered("P0: Fatura, stok fişi, sevk ve sayım için iptal/ters kayıt standardını kurun.")
numbered("P0: SQL Server integration test projesini oluşturun; transaction rollback ve eşzamanlı stok kontrolünü test edin.")
numbered("P1: Ortak InventoryBalanceService kurun; ürün + varyant + depo + rezervasyon hesabını bütün servislerde kullanın.")
numbered("P1: Tahsilat/tediye posting servisini cari ve kasa/banka hareketleriyle atomik hale getirin.")
numbered("P1: AuditLog üretimini merkezi SaveChanges interceptor ile otomatikleştirin.")
numbered("P1: Belge numarası üretimini NumberSequence üzerinden fatura, fiş, sayım, sevk ve tahsilata genişletin.")
numbered("P1: Migration’ı boş/test SQL Server veritabanına uygulayıp seed ve foreign key kontrollerini çalıştırın.")
numbered("P2: Application/Infrastructure katman ayrımını tamamlayın; ardından modül ekranlarına geçin.")

doc.add_heading("13. Canlıya geçmeden önce kabul kriterleri", level=1)
for item in [
    "Her stok değişimi kaynak belge ve hareket satırı üzerinden izlenebiliyor.",
    "Ürün kartı stok miktarı ile depo hareket toplamı arasında fark bulunmuyor.",
    "Onaylanan her belge tekrar onaylanamıyor; idempotency testi geçiyor.",
    "İptal işlemi silme yapmıyor, ters kayıt oluşturuyor.",
    "Satış faturası aynı transaction’da stok ve cariyi etkiliyor.",
    "Alış faturası aynı transaction’da stok ve cariyi etkiliyor.",
    "Tahsilat/tediye aynı transaction’da cari ve kasa/bankayı etkiliyor.",
    "Sayım sonucu sistem/sayılan/fark bilgilerini ve +/− hareketi saklıyor.",
    "Aynı barkod ve stok kodu yarış koşulunda dahi ikinci kez kaydedilemiyor.",
    "Yönetici ve personel izinleri otomatik testlerle doğrulanıyor.",
    "Migration temiz SQL Server 2022 Express üzerinde baştan sona uygulanıyor.",
    "Günlük SQL yedeği ve geri yükleme provası başarıyla tamamlanıyor.",
]:
    bullet("☐ " + item)

doc.add_heading("14. Sonuç", level=1)
doc.add_paragraph(
    "ŞahinSoft sistemi için genişleyebilir bir ön muhasebe veri modeli oluşturulmuştur. Özellikle belge satırlarının stok ve cari hareketlerine foreign key ile bağlanması, GUID kimlikleri, benzersiz indeksler, parametrik stok kuralları ve transaction kullanan posting servisleri doğru yöndedir."
)
doc.add_paragraph(
    "Ancak altyapı denetiminin sonucu “canlıya hazır” değildir. Öncelik; stok bakiyesini tek kaynaktan üretmek, authentication hattını tamamlamak, ters kayıt standardını kurmak ve otomatik testlerle bütün bağlantıları kanıtlamaktır. Bu dört konu tamamlandıktan sonra ekran geliştirmesine güvenle hız verilebilir."
)

doc.core_properties.title = "ŞahinSoft Teknik Altyapı Denetim Raporu"
doc.core_properties.subject = "Ön muhasebe, stok, cari ve teklif sistemi altyapı denetimi"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
